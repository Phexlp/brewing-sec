import re
import io
import hashlib
from typing import List, Dict, Tuple, Any
import fitz  # PyMuPDF
import pdfplumber  # pdfplumber
import docx  # python-docx
import spacy
from spacy.pipeline import EntityRuler

from app.models import ParsedCV, ExtractedEntity

# 1. Build spaCy NLP pipeline with custom EntityRuler for Cybersecurity Entities
nlp = spacy.blank("en")
ruler = nlp.add_pipe("entity_ruler")

# -------------------------------------------------------------------
# CERTIFICATIONS: Only match real cybersecurity certifications.
# Each key is the exact lowercase form we search for in text.
# -------------------------------------------------------------------
CERTIFICATIONS_MAP = {
    "oscp": "Offensive Security Certified Professional (OSCP)",
    "osce": "Offensive Security Certified Expert (OSCE)",
    "oswe": "Offensive Security Web Expert (OSWE)",
    "ceh": "Certified Ethical Hacker (CEH)",
    "cissp": "Certified Information Systems Security Professional (CISSP)",
    "gcfa": "GIAC Certified Forensic Analyst (GCFA)",
    "gcih": "GIAC Certified Incident Handler (GCIH)",
    "grem": "GIAC Reverse Engineering Malware (GREM)",
    "gpen": "GIAC Penetration Tester (GPEN)",
    "gxpn": "GIAC Exploit Researcher (GXPN)",
    "gnfa": "GIAC Network Forensic Analyst (GNFA)",
    "gsec": "GIAC Security Essentials (GSEC)",
    "splunk core certified": "Splunk Core Certified User",
    "cysa+": "CompTIA CySA+",
    "security+": "CompTIA Security+",
    "sec+": "CompTIA Security+",
    "comptia security+": "CompTIA Security+",
    "pentest+": "CompTIA PenTest+",
    "ejpt": "eLearnSecurity Junior Penetration Tester (eJPT)",
    "pnpt": "Practical Network Penetration Tester (PNPT)",
    "cisa": "Certified Information Systems Auditor (CISA)",
    "cism": "Certified Information Security Manager (CISM)",
    "casp+": "CompTIA Advanced Security Practitioner (CASP+)",
    "ecsa": "EC-Council Certified Security Analyst (ECSA)",
    "chfi": "Computer Hacking Forensic Investigator (CHFI)",
    "ccna": "Cisco Certified Network Associate (CCNA)",
    "ccnp": "Cisco Certified Network Professional (CCNP)",
    "ccie": "Cisco Certified Internetwork Expert (CCIE)",
    "aws certified": "AWS Certified Security Specialty",
    "az-500": "Microsoft Azure Security Engineer (AZ-500)",
    "sc-200": "Microsoft Security Operations Analyst (SC-200)"
}

# -------------------------------------------------------------------
# CYBER_KEYWORDS: domain-specific terms the parser searches for.
# These are real, specific technical terms — NOT generic words.
# A CV that doesn't mention these terms gets zero domain scores.
# -------------------------------------------------------------------
CYBER_KEYWORDS = {
    "web_security": [
        "sql injection", "sqli", "xss", "cross-site scripting",
        "ssrf", "csrf", "idor", "bola",
        "burp suite", "burpsuite", "owasp", "owasp top 10",
        "command injection", "ssti", "template injection",
        "broken authentication", "session hijacking",
        "api security", "api testing", "postman",
        "google dorking", "reconnaissance", "subdomain enumeration",
        "race condition", "web application firewall", "waf bypass",
        "deserialization", "insecure direct object",
        "pentesting", "penetration testing", "web pentesting",
        "zap", "nikto", "dirb", "gobuster", "ffuf"
    ],
    "network_security": [
        "active directory", "bloodhound", "powerview",
        "kerberoasting", "as-rep roasting", "mimikatz",
        "nmap", "wireshark", "tcpdump", "tshark",
        "pivoting", "chisel", "ligolo-ng", "socks proxy",
        "metasploit", "msfconsole", "meterpreter",
        "subnetting", "cidr", "tcp/ip", "osi model",
        "nat", "port forwarding", "ssh tunneling",
        "firewall", "iptables", "packet analysis",
        "linux", "bash", "chmod", "chown",
        "kali linux", "parrot os",
        "privilege escalation", "privesc",
        "lateral movement", "pass the hash"
    ],
    "dfir": [
        "volatility", "memory forensics", "memory dump",
        "autopsy", "ftk imager", "ftk", "encase",
        "disk forensics", "data carving", "file recovery",
        "shimcache", "amcache", "prefetch",
        "registry analysis", "regripper",
        "incident response", "digital forensics",
        "chain of custody", "evidence collection",
        "timeline reconstruction", "super timeline",
        "steganography", "metadata analysis", "exiftool",
        "linux forensics", "windows forensics",
        "pcap analysis", "network forensics",
        "sleuthkit", "scalpel", "foremost",
        "alternate data streams",
        "ransomware forensics", "malware forensics"
    ],
    "soc_siem": [
        "splunk", "spl", "splunk search",
        "elastic", "elasticsearch", "kibana", "kql",
        "sysmon", "sysmon event", "event id",
        "snort", "suricata", "nids",
        "soar", "qradar", "arcsight",
        "log analysis", "siem", "log correlation",
        "phishing analysis", "email analysis",
        "edr", "crowdstrike", "defender for endpoint",
        "siem alert", "triage", "alert investigation",
        "syslog", "cef", "leef",
        "apache log", "access log", "web server log",
        "log pivoting", "log types"
    ],
    "threat_hunting": [
        "mitre att&ck", "mitre attack", "att&ck",
        "sigma rules", "sigma", "hayabusa",
        "cobalt strike", "sliver", "c2 beacon",
        "yara", "yara rules",
        "threat hunting", "threat intelligence",
        "powershell hunting", "script block logging",
        "persistence", "scheduled tasks", "wmi persistence",
        "opencti", "misp", "threat feeds",
        "github recon", "osint",
        "threat modelling", "threat modeling",
        "stride", "dread",
        "container security", "docker forensics",
        "privilege escalation", "privesc",
        "apt", "advanced persistent threat"
    ],
    "malware_re": [
        "ghidra", "ida pro", "ida free",
        "x64dbg", "ollydbg", "windbg",
        "assembly", "x86", "x64", "disassembly",
        "pe structure", "pe header", "pestudio",
        "dnspy", "ilspy", ".net decompilation",
        "speakeasy", "unpacking", "upx",
        "reverse engineering", "malware analysis",
        "static analysis", "dynamic analysis",
        "anti-debugging", "anti-analysis", "anti-vm",
        "process injection", "dll injection", "process hollowing",
        "shellcode", "obfuscation", "packing",
        "sandbox", "cuckoo", "any.run",
        "behavior monitoring", "api monitoring",
        "strings", "file hashing", "virustotal"
    ]
}

# -------------------------------------------------------------------
# JOB TITLES: Only cybersecurity-specific titles.
# -------------------------------------------------------------------
JOB_TITLES = [
    "soc analyst", "security analyst", "security engineer",
    "penetration tester", "pentester", "red teamer",
    "dfir specialist", "incident responder", "forensic analyst",
    "threat hunter", "threat analyst",
    "malware analyst", "reverse engineer",
    "security consultant", "security architect",
    "vulnerability analyst", "blue teamer",
    "cybersecurity analyst", "information security analyst",
    "security operations", "devsecops"
]

# Register patterns into spaCy EntityRuler
patterns = []
for cert_key in CERTIFICATIONS_MAP.keys():
    patterns.append({"label": "CERT", "pattern": [{"LOWER": token} for token in cert_key.split()]})

for domain, kw_list in CYBER_KEYWORDS.items():
    for kw in kw_list:
        patterns.append({"label": "SKILL", "pattern": [{"LOWER": token} for token in kw.split()]})

for title in JOB_TITLES:
    patterns.append({"label": "JOB_TITLE", "pattern": [{"LOWER": token} for token in title.split()]})

ruler.add_patterns(patterns)


def extract_text_from_pdf(content: bytes) -> str:
    """Extract raw text using pdfplumber with PyMuPDF (fitz) fallback."""
    text_chunks = []
    # Primary: pdfplumber
    try:
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page in pdf.pages:
                txt = page.extract_text()
                if txt:
                    text_chunks.append(txt)
    except Exception as e:
        print(f"pdfplumber extraction note: {e}")

    if not text_chunks:
        # Fallback: PyMuPDF fitz
        try:
            with fitz.open(stream=content, filetype="pdf") as doc:
                for page in doc:
                    text_chunks.append(page.get_text("text"))
        except Exception as e:
            print(f"PyMuPDF extraction note: {e}")

    return "\n".join(text_chunks)


def extract_text_from_docx(content: bytes) -> str:
    """Extract raw text from DOCX using python-docx."""
    doc = docx.Document(io.BytesIO(content))
    full_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    full_text.append(cell.text)
    return "\n".join(full_text)


def parse_experience_years(text: str) -> float:
    """Extract total years of experience using regular expressions.
    Returns 0.0 if no experience indicators are found (NOT a default).
    """
    text_lower = text.lower()

    # Pattern 1: "X years of experience" / "X+ years"
    exp_matches = re.findall(r'(\d+(?:\.\d+)?)\s*\+?\s*(?:years|yrs|year)\s*(?:of)?\s*(?:experience|exp)?', text_lower)
    if exp_matches:
        years = [float(m) for m in exp_matches]
        return max(years)

    # Pattern 2: Date ranges like "2019 - 2023" or "2020 - present"
    date_ranges = re.findall(r'(20\d{2})\s*(?:-|to|–)\s*(20\d{2}|present|current)', text_lower)
    total_years = 0.0
    for start, end in date_ranges:
        start_yr = float(start)
        end_yr = 2026.0 if end in ["present", "current"] else float(end)
        diff = max(0.5, end_yr - start_yr)
        total_years += diff

    # Return 0.0 if nothing found — do NOT fabricate experience
    return round(total_years, 1)


def parse_cv_nlp(filename: str, content: bytes) -> ParsedCV:
    """NLP parsing entrypoint powered by spaCy EntityRuler, PyMuPDF/pdfplumber & python-docx.

    CRITICAL: This function must NOT hallucinate. If a keyword is not in the text,
    it is not extracted. If no cybersecurity terms are found, all domain scores are 0.
    """
    ext = filename.lower().split('.')[-1]
    if ext == 'pdf':
        raw_text = extract_text_from_pdf(content)
    elif ext in ['docx', 'doc']:
        raw_text = extract_text_from_docx(content)
    else:
        raw_text = content.decode('utf-8', errors='ignore')

    text_lower = raw_text.lower()
    cv_hash = hashlib.sha256(raw_text.encode('utf-8')).hexdigest()[:16]

    # Process text through spaCy NLP pipeline
    doc = nlp(raw_text)

    found_certs = set()
    found_skills = set()
    found_titles = set()
    entities: List[ExtractedEntity] = []

    # Process entities extracted by spaCy EntityRuler
    for ent in doc.ents:
        ent_text = ent.text.strip()
        ent_lower = ent_text.lower()
        if ent.label_ == "CERT":
            full_cert = CERTIFICATIONS_MAP.get(ent_lower, ent_text.upper())
            found_certs.add(full_cert)
            entities.append(ExtractedEntity(text=full_cert, label="CERT", category="certification"))
        elif ent.label_ == "SKILL":
            found_skills.add(ent_text.title())
            entities.append(ExtractedEntity(text=ent_text.title(), label="SKILL", category="cyber_skill"))
        elif ent.label_ == "JOB_TITLE":
            found_titles.add(ent_text.title())
            entities.append(ExtractedEntity(text=ent_text.title(), label="JOB_TITLE", category="experience"))

    # Regex keyword fallback — only count ACTUAL matches in the text
    domain_counts = {dom: 0 for dom in CYBER_KEYWORDS.keys()}
    for domain, kw_list in CYBER_KEYWORDS.items():
        for kw in kw_list:
            # Use word boundary matching to avoid false positives
            pattern = r'\b' + re.escape(kw) + r'\b'
            if re.search(pattern, text_lower):
                found_skills.add(kw.title())
                domain_counts[domain] += 1

    # Certification regex fallback
    for cert_key, cert_full in CERTIFICATIONS_MAP.items():
        pattern = r'\b' + re.escape(cert_key) + r'\b'
        if re.search(pattern, text_lower):
            found_certs.add(cert_full)

    # Job title regex fallback
    for title in JOB_TITLES:
        pattern = r'\b' + re.escape(title) + r'\b'
        if re.search(pattern, text_lower):
            found_titles.add(title.title())

    # DO NOT add a default job title if none found.
    # An accountant's CV should show zero cybersecurity titles.

    exp_years = parse_experience_years(raw_text)
    if exp_years > 0:
        entities.append(ExtractedEntity(text=f"{exp_years} Years Experience", label="EXP", category="experience"))

    return ParsedCV(
        raw_text_length=len(raw_text),
        filename=filename,
        skills=sorted(list(found_skills)),
        certifications=sorted(list(found_certs)),
        job_titles=sorted(list(found_titles)),
        experience_years=exp_years,
        detected_domains=domain_counts,
        entities=entities,
        cv_hash=cv_hash
    )
