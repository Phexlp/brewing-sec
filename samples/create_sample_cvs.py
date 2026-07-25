import fitz  # PyMuPDF
import docx

def create_sample_pdf(filepath: str):
    doc = fitz.open()
    page = doc.new_page()
    
    cv_text = """
    ALEXANDER RIVERS
    Senior SOC & Incident Response Specialist
    Email: alex.rivers@secmail.io | Mobile: +1-555-019-2834
    
    PROFESSIONAL SUMMARY
    Dedicated Security Analyst with 4.5 years of experience in enterprise SOC monitoring, alert triage, SIEM detection engineering, and digital forensics. Proven track record in hunting threats using Splunk SPL, Sysmon telemetry, and Volatility memory analysis.
    
    CERTIFICATIONS
    - CompTIA Security+ (Sec+)
    - CompTIA CySA+
    - Splunk Core Certified Power User
    - GIAC Certified Incident Handler (GCIH)
    
    TECHNICAL SKILLS & COMPETENCIES
    - SIEM / Logging: Splunk (SPL), Elastic / Kibana (KQL), Sysmon, Windows Event Logs (4624, 4688, 4104)
    - Forensics & DFIR: Volatility 3 RAM analysis, Autopsy Disk Forensics, Registry Hive Parsing (Shimcache, Amcache), Wireshark PCAP
    - Threat Hunting: MITRE ATT&CK Mapping, YARA rules, Sigma detection rules, PowerShell Script Block Logging
    - Security Tools: Burp Suite, Nmap, Snort, Suricata, CrowdStrike Falcon EDR, Python Scripting
    
    PROFESSIONAL EXPERIENCE
    Cyber Defense Center - Senior SOC Analyst (2022 - Present)
    - Triaged over 500 high-severity security incidents using Splunk SIEM and automated Shuffle SOAR playbooks.
    - Authored 30+ Sigma detection rules for detecting Cobalt Strike beaconing and memory injection techniques.
    - Conducted memory forensic analysis using Volatility 3 on infected domain endpoints.
    
    Defensive Systems - Junior Security Analyst (2020 - 2022)
    - Monitored network traffic using Wireshark and Suricata NIDS.
    - Analyzed malicious phishing email attachments and executed safe detonation in Cuckoo Sandbox.
    """
    
    page.insert_text((50, 50), cv_text, fontsize=10)
    doc.save(filepath)
    doc.close()
    print(f"Created PDF: {filepath}")

def create_sample_docx(filepath: str):
    doc = docx.Document()
    doc.add_heading("MARCUS VANCE", level=0)
    doc.add_paragraph("Offensive Security & Penetration Testing Specialist | OSCP, eJPT")
    doc.add_paragraph("Experience: 3 years in Web Application & Network Penetration Testing")
    
    doc.add_heading("Certifications", level=1)
    doc.add_paragraph("- Offensive Security Certified Professional (OSCP)")
    doc.add_paragraph("- eLearnSecurity Junior Penetration Tester (eJPT)")
    
    doc.add_heading("Technical Expertise", level=1)
    doc.add_paragraph("Web Security: SQL Injection (SQLi), Cross-Site Scripting (XSS), SSRF, OAuth 2.0, IDOR, Burp Suite Pro, OWASP Top 10")
    doc.add_paragraph("Network & Active Directory: BloodHound, PowerView, Kerberoasting, AS-REP Roasting, Mimikatz, Nmap, Chisel Pivoting, Metasploit")
    doc.add_paragraph("Reverse Engineering Basics: Ghidra static analysis, x64dbg debugging, PEStructure inspection")
    
    doc.add_heading("Work History", level=1)
    doc.add_paragraph("Red Team Consultant (2021 - Present) at CyberStrike Ops")
    doc.add_paragraph("- Executed external and internal network penetration tests against Fortune 500 Active Directory environments.")
    doc.add_paragraph("- Discovered zero-day SSRF and Deserialization vulnerabilities in corporate web portals.")
    
    doc.save(filepath)
    print(f"Created DOCX: {filepath}")

if __name__ == "__main__":
    create_sample_pdf("C:/Users/Phexl/.gemini/antigravity/scratch/career-mapper/samples/sample_soc_analyst.pdf")
    create_sample_docx("C:/Users/Phexl/.gemini/antigravity/scratch/career-mapper/samples/sample_pentester.docx")
